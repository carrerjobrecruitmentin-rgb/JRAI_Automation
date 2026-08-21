import os
import re
import json
import tempfile
from typing import Optional, Dict, Any
from common.config import settings
from common.logger import log
from services.cv_parser.schemas import ParsedCVResponse

# Optional OCR & Document imports
try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

import requests
try:
    from groq import Groq
except ImportError:
    Groq = None


class CVParserService:
    """
    Production CV & Resume Parser Service.
    Extracts text from PDF/DOCX/Images and normalizes into rich structured profiles using AI.
    """

    @classmethod
    def extract_text_from_bytes(cls, file_bytes: bytes, filename: str) -> str:
        """
        Extracts raw text from PDF, DOCX, TXT, or Image bytes.
        """
        filename_lower = filename.lower()
        extracted_text = ""

        # 1. Plain Text
        if filename_lower.endswith(".txt"):
            try:
                return file_bytes.decode("utf-8", errors="ignore")
            except Exception as e:
                log.warning(f"Error decoding txt file: {e}")

        # 2. DOCX Files
        elif filename_lower.endswith(".docx") and docx:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                doc = docx.Document(tmp_path)
                full_text = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            full_text.append(" | ".join(row_text))
                os.remove(tmp_path)
                return "\n".join(full_text)
            except Exception as e:
                log.warning(f"DOCX extraction failed: {e}")

        # 3. PDF Files (pdfplumber -> pypdf -> OCR fallback)
        elif filename_lower.endswith(".pdf"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            # Method A: pdfplumber
            if pdfplumber:
                try:
                    with pdfplumber.open(tmp_path) as pdf:
                        pages_text = []
                        for page in pdf.pages:
                            text = page.extract_text()
                            if text:
                                pages_text.append(text)
                        if pages_text:
                            extracted_text = "\n".join(pages_text)
                except Exception as e:
                    log.warning(f"pdfplumber extraction failed: {e}")

            # Method B: pypdf fallback
            if not extracted_text.strip() and pypdf:
                try:
                    reader = pypdf.PdfReader(tmp_path)
                    pages_text = [p.extract_text() for p in reader.pages if p.extract_text()]
                    if pages_text:
                        extracted_text = "\n".join(pages_text)
                except Exception as e:
                    log.warning(f"pypdf extraction failed: {e}")

            if os.path.exists(tmp_path):
                os.remove(tmp_path)

        # 4. Images (Tesseract OCR)
        elif any(filename_lower.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".webp"]) and pytesseract and Image:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".img") as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                image = Image.open(tmp_path)
                extracted_text = pytesseract.image_to_string(image)
                os.remove(tmp_path)
            except Exception as e:
                log.warning(f"Tesseract OCR image extraction failed: {e}")

        return extracted_text.strip()

    @classmethod
    def parse_with_ai(cls, raw_text: str, target_role: Optional[str] = None) -> ParsedCVResponse:
        """
        Takes raw resume text and uses Gemini Flash (primary) or Groq LLaMA (fallback) to produce structured JSON.
        """
        if not raw_text or len(raw_text.strip()) < 20:
            return ParsedCVResponse(
                success=False,
                message="Text is too short or empty to parse as a resume",
                raw_text_preview=raw_text[:200]
            )

        prompt = f"""
You are an expert AI Resume Parser. Extract all professional details from this candidate's resume and return STRICT, VALID JSON ONLY (no markdown formatting, no code blocks, just raw JSON).

Schema:
{{
  "full_name": "Full Name",
  "email": "candidate@email.com",
  "phone": "+91 9876543210",
  "location": "City, State, Country",
  "current_role": "Current or Latest Job Title",
  "total_experience_years": 3.5,
  "skills": ["Skill 1", "Skill 2", "Skill 3"],
  "education": [
    {{
      "degree": "B.Tech Computer Science",
      "institution": "University Name",
      "graduation_year": "2024",
      "percentage_or_cgpa": "8.5 CGPA"
    }}
  ],
  "experience": [
    {{
      "job_title": "Software Engineer",
      "company_name": "Company Name",
      "location": "City",
      "start_date": "2022",
      "end_date": "Present",
      "is_current": true,
      "responsibilities": ["Built REST APIs", "Optimized database queries"]
    }}
  ],
  "certifications": ["AWS Certified", "Python Professional"],
  "languages": ["English", "Hindi"],
  "summary": "2-3 sentence executive professional summary of candidate"
}}

Resume Text:
\"\"\"
{raw_text[:8000]}
\"\"\"
"""

        # Tier 1: Gemini 2.5 Flash
        if settings.GEMINI_API_KEY:
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{settings.GEMINI_MODEL}:generateContent?key={settings.GEMINI_API_KEY}"
                payload = {
                    "contents": [{"parts": [{"text": prompt}]}],
                    "generationConfig": {"temperature": 0.1}
                }
                resp = requests.post(url, json=payload, timeout=25)
                if resp.status_code == 200:
                    data = resp.json()
                    ai_text = data["candidates"][0]["content"]["parts"][0]["text"]
                    cleaned_json = cls._clean_json_string(ai_text)
                    parsed_dict = json.loads(cleaned_json)
                    parsed_dict["success"] = True
                    parsed_dict["extracted_by"] = "gemini_2.5_flash"
                    parsed_dict["raw_text_preview"] = raw_text[:300]
                    return ParsedCVResponse(**parsed_dict)
            except Exception as e:
                log.warning(f"Gemini CV parser failed, attempting Groq fallback: {e}")

        # Tier 2: Groq LLaMA 3.3 70B Versatile
        if settings.GROQ_API_KEY and Groq:
            try:
                client = Groq(api_key=settings.GROQ_API_KEY)
                chat_completion = client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "You are a professional resume parser returning valid JSON only."},
                        {"role": "user", "content": prompt}
                    ],
                    model=settings.GROQ_MODEL,
                    temperature=0.1,
                    response_format={"type": "json_object"}
                )
                ai_text = chat_completion.choices[0].message.content
                parsed_dict = json.loads(ai_text)
                parsed_dict["success"] = True
                parsed_dict["extracted_by"] = "groq_llama_3.3"
                parsed_dict["raw_text_preview"] = raw_text[:300]
                return ParsedCVResponse(**parsed_dict)
            except Exception as e:
                log.warning(f"Groq CV parser failed: {e}")

        # Tier 3: Regex / Heuristic Fallback
        return cls._heuristic_fallback(raw_text)

    @classmethod
    def _clean_json_string(cls, text: str) -> str:
        text = re.sub(r"^```json\s*", "", text.strip(), flags=re.MULTILINE)
        text = re.sub(r"^```\s*$", "", text.strip(), flags=re.MULTILINE)
        return text.strip()

    @classmethod
    def _heuristic_fallback(cls, text: str) -> ParsedCVResponse:
        emails = re.findall(r"[\w\.-]+@[\w\.-]+\.\w+", text)
        phones = re.findall(r"(?:\+91[\-\s]?)?[6-9]\d{9}", text)

        email = emails[0] if emails else None
        phone = phones[0] if phones else None

        # Extract common tech skills
        tech_keywords = [
            "Python", "JavaScript", "PHP", "Laravel", "React", "Node.js", "MySQL",
            "PostgreSQL", "Docker", "AWS", "HTML", "CSS", "Tailwind", "Java", "C++",
            "Git", "REST API", "Figma", "DevOps", "Linux", "Kubernetes"
        ]
        found_skills = [k for k in tech_keywords if re.search(rf"\b{re.escape(k)}\b", text, re.IGNORECASE)]

        return ParsedCVResponse(
            success=True,
            message="Parsed using heuristic extraction fallback",
            email=email,
            phone=phone,
            skills=found_skills,
            extracted_by="heuristic_fallback",
            raw_text_preview=text[:300]
        )
